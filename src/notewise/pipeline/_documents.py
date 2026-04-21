"""Document rendering helpers for notes output formats."""

from __future__ import annotations

import re
import unicodedata
from collections.abc import Callable
from html import escape
from io import BytesIO
from pathlib import Path

from notewise._constants import (
    CHAPTER_BUNDLE_SEPARATOR,
    DEFAULT_NOTES_OUTPUT_FORMAT,
    DEFAULT_RENDERED_HTML_LANG,
    DEFAULT_RENDERED_HTML_STYLES,
    DEFAULT_TARGET_LANGUAGE,
    DOCX_BODY_FONT_NAME,
    DOCX_BODY_FONT_SIZE_PT,
    DOCX_BODY_SPACE_AFTER_PT,
    DOCX_HEADING_FONT_NAME,
    DOCX_HEADING_ONE_FONT_SIZE_PT,
    DOCX_HEADING_SPACE_AFTER_PT,
    DOCX_HEADING_SPACE_BEFORE_PT,
    DOCX_HEADING_THREE_FONT_SIZE_PT,
    DOCX_HEADING_TWO_FONT_SIZE_PT,
    DOCX_SECTION_MARGIN_INCHES,
    DOCX_TITLE_FONT_SIZE_PT,
    HTML_LANGUAGE_ALIASES,
    MARKDOWN_RENDER_EXTENSIONS,
    NOTES_OUTPUT_EXTENSIONS,
    OUTPUT_FORMAT_SEPARATOR,
    PDF_UNSUPPORTED_UNICODE_ERROR,
    SUPPORTED_NOTES_OUTPUT_FORMATS,
)
from notewise.errors import ValidationError


DocumentRenderer = Callable[[str, str, Path, str | None], None]
_LIST_ITEM_RE = re.compile(r"^(?P<indent>\s*)(?:[-*+]\s+|\d+\.\s+)")
_CODE_BLOCK_RE = re.compile(
    r"<pre><code(?:\s+class=\"[^\"]*\")?>(?P<code>.*?)</code></pre>", re.DOTALL
)
_LANGUAGE_CODE_RE = re.compile(r"^[A-Za-z]{2,3}(?:-[A-Za-z0-9]{2,8})*$")
_PDF_CHARACTER_TRANSLATIONS = str.maketrans(
    {
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u2026": "...",
        "\u00a0": " ",
    }
)


def normalize_output_format(output_format: str | None) -> str:
    """Normalize and validate the requested notes output format."""
    candidate = (output_format or DEFAULT_NOTES_OUTPUT_FORMAT).strip().lower()
    if candidate not in SUPPORTED_NOTES_OUTPUT_FORMATS:
        supported = ", ".join(SUPPORTED_NOTES_OUTPUT_FORMATS)
        raise ValidationError(
            f"Unsupported output format '{output_format}'. Choose one of: {supported}."
        )
    return candidate


def normalize_output_formats(
    output_format_value: str | list[str] | tuple[str, ...] | None,
) -> list[str]:
    """Parse, validate, and deduplicate one or more requested formats."""
    if output_format_value is None:
        raw_values = [DEFAULT_NOTES_OUTPUT_FORMAT]
    elif isinstance(output_format_value, str):
        raw_values = output_format_value.split(OUTPUT_FORMAT_SEPARATOR)
    else:
        raw_values = list(output_format_value)

    normalized_formats: list[str] = []
    for raw_value in raw_values:
        candidate = normalize_output_format(raw_value)
        if candidate not in normalized_formats:
            normalized_formats.append(candidate)
    return normalized_formats


def get_output_extension(output_format: str | None) -> str:
    """Return the canonical file extension for the selected format."""
    return NOTES_OUTPUT_EXTENSIONS[normalize_output_format(output_format)]


def build_chapter_bundle(title: str, chapter_notes: list[str]) -> str:
    """Bundle generated chapter notes into one Markdown document."""
    cleaned_notes = [notes.strip() for notes in chapter_notes if notes.strip()]
    if not cleaned_notes:
        return f"# {title}\n"
    return f"# {title}\n\n{CHAPTER_BUNDLE_SEPARATOR.join(cleaned_notes)}\n"


def render_notes_document(
    markdown_text: str,
    title: str,
    output_path: Path,
    output_format: str | None,
    target_language: str | None = None,
) -> Path:
    """Render a Markdown study document to the requested file format."""
    normalized_format = normalize_output_format(output_format)
    renderer = _DOCUMENT_RENDERERS[normalized_format]
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        renderer(markdown_text, title, output_path, target_language)
        return output_path
    except ValidationError as error:
        if normalized_format != "pdf" or not str(error).startswith(
            "PDF output currently supports Latin-script text only."
        ):
            raise

        fallback_path = output_path.with_suffix(NOTES_OUTPUT_EXTENSIONS["md"])
        _write_markdown(markdown_text, title, fallback_path, target_language)
        return fallback_path


def render_notes_documents(
    markdown_text: str,
    title: str,
    output_targets: dict[str, Path],
    target_language: str | None = None,
) -> dict[str, Path]:
    """Render a Markdown study document to multiple target formats."""
    rendered_targets: dict[str, Path] = {}
    for output_format, output_path in output_targets.items():
        rendered_targets[output_format] = render_notes_document(
            markdown_text,
            title,
            output_path,
            output_format,
            target_language=target_language,
        )
    return rendered_targets


def _normalize_markdown_blocks(markdown_text: str) -> str:
    normalized_lines: list[str] = []
    previous_line = ""

    for line in markdown_text.splitlines():
        is_list_item = bool(_LIST_ITEM_RE.match(line))
        previous_is_list_item = bool(_LIST_ITEM_RE.match(previous_line))
        previous_is_content = bool(previous_line.strip())
        previous_is_heading = previous_line.lstrip().startswith("#")

        if (
            is_list_item
            and not previous_is_list_item
            and previous_is_content
            and not previous_is_heading
            and normalized_lines
            and normalized_lines[-1] != ""
        ):
            normalized_lines.append("")

        normalized_lines.append(line)
        previous_line = line

    return "\n".join(normalized_lines)


def _markdown_to_html(markdown_text: str) -> str:
    import markdown as markdown_lib

    rendered_html = markdown_lib.markdown(
        _normalize_markdown_blocks(markdown_text),
        extensions=list(MARKDOWN_RENDER_EXTENSIONS),
    )
    return _normalize_rendered_html(rendered_html)


def _normalize_rendered_html(body_html: str) -> str:
    def _replace_code_block(match: re.Match[str]) -> str:
        code_html = match.group("code")
        return f'<pre class="code-block"><code>{code_html}</code></pre>'

    return _CODE_BLOCK_RE.sub(_replace_code_block, body_html)


def _normalize_pdf_markdown(
    markdown_text: str,
    target_language: str | None = None,
) -> str:
    translated = markdown_text.translate(_PDF_CHARACTER_TRANSLATIONS)
    normalized = unicodedata.normalize("NFKC", translated)
    try:
        normalized.encode("latin-1")
    except UnicodeEncodeError as error:
        language_label = (target_language or DEFAULT_TARGET_LANGUAGE).strip()
        raise ValidationError(
            PDF_UNSUPPORTED_UNICODE_ERROR.format(target_language=language_label)
        ) from error
    return normalized


def _resolve_html_lang(target_language: str | None) -> str:
    candidate = (target_language or DEFAULT_TARGET_LANGUAGE).strip()
    if not candidate:
        return DEFAULT_RENDERED_HTML_LANG

    if _LANGUAGE_CODE_RE.match(candidate):
        return candidate.replace("_", "-")

    return HTML_LANGUAGE_ALIASES.get(
        candidate.casefold(),
        DEFAULT_RENDERED_HTML_LANG,
    )


def _build_html_document(
    title: str,
    body_html: str,
    styles: str = DEFAULT_RENDERED_HTML_STYLES,
    lang: str = DEFAULT_RENDERED_HTML_LANG,
) -> str:
    escaped_title = escape(title)
    return (
        "<!DOCTYPE html>\n"
        f'<html lang="{escape(lang)}">\n'
        "<head>\n"
        '  <meta charset="utf-8" />\n'
        '  <meta name="viewport" content="width=device-width, initial-scale=1" />\n'
        f"  <title>{escaped_title}</title>\n"
        "  <style>\n"
        f"{styles}\n"
        "  </style>\n"
        "</head>\n"
        "<body>\n"
        '  <main class="document">\n'
        f"{body_html}\n"
        "  </main>\n"
        "</body>\n"
        "</html>\n"
    )


def _write_markdown(
    markdown_text: str,
    _title: str,
    output_path: Path,
    _target_language: str | None = None,
) -> None:
    output_path.write_text(markdown_text, encoding="utf-8")


def _write_html(
    markdown_text: str,
    title: str,
    output_path: Path,
    target_language: str | None = None,
) -> None:
    html_document = _build_html_document(
        title,
        _markdown_to_html(markdown_text),
        lang=_resolve_html_lang(target_language),
    )
    output_path.write_text(html_document, encoding="utf-8")


def _write_pdf(
    markdown_text: str,
    title: str,
    output_path: Path,
    target_language: str | None = None,
) -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_title(title)
    pdf.set_author("notewise")
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.set_margins(16, 16, 16)
    pdf.add_page()
    pdf.write_html(
        _markdown_to_html(
            _normalize_pdf_markdown(markdown_text, target_language=target_language)
        ),
        font_family="Times",
    )
    pdf.output(str(output_path))


def _write_docx(
    markdown_text: str,
    title: str,
    output_path: Path,
    _target_language: str | None = None,
) -> None:
    from docx import Document
    from docx.shared import Inches, Pt
    from html2docx import html2docx

    document_buffer = html2docx(_markdown_to_html(markdown_text), title=title)
    document_buffer.seek(0)
    document = Document(document_buffer)

    for section in document.sections:
        section.top_margin = Inches(DOCX_SECTION_MARGIN_INCHES)
        section.bottom_margin = Inches(DOCX_SECTION_MARGIN_INCHES)
        section.left_margin = Inches(DOCX_SECTION_MARGIN_INCHES)
        section.right_margin = Inches(DOCX_SECTION_MARGIN_INCHES)

    normal_style = document.styles["Normal"]
    normal_style.font.name = DOCX_BODY_FONT_NAME
    normal_style.font.size = Pt(DOCX_BODY_FONT_SIZE_PT)

    for style_name, font_size in (
        ("Title", DOCX_TITLE_FONT_SIZE_PT),
        ("Heading 1", DOCX_HEADING_ONE_FONT_SIZE_PT),
        ("Heading 2", DOCX_HEADING_TWO_FONT_SIZE_PT),
        ("Heading 3", DOCX_HEADING_THREE_FONT_SIZE_PT),
    ):
        style = document.styles[style_name]
        style.font.name = DOCX_HEADING_FONT_NAME
        style.font.size = Pt(font_size)
        style.paragraph_format.space_before = Pt(DOCX_HEADING_SPACE_BEFORE_PT)
        style.paragraph_format.space_after = Pt(DOCX_HEADING_SPACE_AFTER_PT)

    def _apply_paragraph_style(paragraph) -> None:
        if paragraph.text.strip():
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(DOCX_BODY_SPACE_AFTER_PT)
            paragraph.paragraph_format.line_spacing = 1.1
        else:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

    for paragraph in document.paragraphs:
        _apply_paragraph_style(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _apply_paragraph_style(paragraph)

    styled_buffer = BytesIO()
    document.save(styled_buffer)
    output_path.write_bytes(styled_buffer.getvalue())


_DOCUMENT_RENDERERS: dict[str, DocumentRenderer] = {
    "md": _write_markdown,
    "html": _write_html,
    "pdf": _write_pdf,
    "docx": _write_docx,
}
